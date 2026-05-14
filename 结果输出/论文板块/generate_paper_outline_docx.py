# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\math\2026数模题目\B题")
OUT = ROOT / "B题_solution_outputs"
EXP = OUT / "实验板块"
FIG = OUT / "figures2.1"
PAPER = OUT / "论文板块"
DOCX_PATH = PAPER / "B题_论文大纲与对应草稿.docx"


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    for style_name in ["Normal", "Body Text"]:
        style = styles[style_name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)
        style.paragraph_format.line_spacing = 1.35
        style.paragraph_format.space_after = Pt(4)

    heading_specs = {
        "Title": (18, True),
        "Heading 1": (15, True),
        "Heading 2": (13, True),
        "Heading 3": (11.5, True),
    }
    for style_name, (size, bold) in heading_specs.items():
        style = styles[style_name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = bold


def add_paragraph(doc: Document, text: str, size=10.5, bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    run = p.add_run("说明：")
    set_run_font(run, bold=True, color=(60, 90, 150))
    run = p.add_run(text)
    set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_run_font(r, size=9.5)
    if widths:
        for row in table.rows:
            for idx, width_cm in enumerate(widths):
                row.cells[idx].width = Cm(width_cm)
    doc.add_paragraph()


def safe_read_csv(name: str) -> pd.DataFrame:
    return pd.read_csv(EXP / name, encoding="utf-8-sig")


def load_metrics():
    pred = safe_read_csv("问题1_第23-29天预测.csv")
    params = safe_read_csv("模型参数与检验.csv")
    p2_main = safe_read_csv("问题2_主效应与交互效应.csv")
    p2_ablation = safe_read_csv("问题2_消融分析.csv")
    p3_key = safe_read_csv("问题3_阈值关键点.csv")
    p3_mono = safe_read_csv("问题3_单调性校验.csv")
    p4 = safe_read_csv("问题4_三区域判定结果.csv")
    exp = pd.read_csv(ROOT / "附件1_模拟实验数据.csv", encoding="utf-8-sig")
    return pred, params, p2_main, p2_ablation, p3_key, p3_mono, p4, exp


def metric_lookup(params: pd.DataFrame) -> dict[str, str]:
    out = {}
    rows = params.fillna("")
    for _, row in rows.iterrows():
        key = str(row.iloc[0]).strip()
        if key in {"r2_log", "rmse_log", "sigma_log", "rmse_m", "mae_m", "mape_m"}:
            out[key] = str(row.iloc[1])
    return out


def add_figure(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        add_note(doc, f"未找到图文件：{path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(14.6))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=9.5)


def build_doc() -> None:
    PAPER.mkdir(parents=True, exist_ok=True)
    pred, params, p2_main, p2_ablation, p3_key, p3_mono, p4, exp = load_metrics()
    m = metric_lookup(params)

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("B题 剩磁法雷击判定预测\n论文大纲与对应草稿")
    set_run_font(r, size=18, bold=True)
    add_paragraph(doc, "资料基础：GitHub 仓库 2026CQUPTmathematical-modeling、附件1模拟实验数据、附件2天气数据、实验板块结果表与 figures2.1 论文图表。", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "用途：本文件用于论文阶段搭建正文骨架，后续可直接进入格式规范、语言润色和图表精排。", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_heading("一、仓库内容回忆与论文主线", level=1)
    add_paragraph(
        doc,
        "已读取的 GitHub 仓库内容显示，项目围绕 B 题“剩磁法雷击判定预测”展开，核心材料包括题面与 PDF 提取文本、原始数据附件、问题二/三/四建模说明、图表需求与数据来源说明、结果输出目录以及完整求解报告。仓库中的图表说明强调图表需要与结果表一致、可复跑、可追溯；论文草稿则已经形成“个体初值锚定 + 环境驱动指数衰减 + 动态阈值 + 三区域判定”的主线。",
    )
    add_paragraph(
        doc,
        f"本地实验结果进一步支持这一主线：附件1共有 {len(exp)} 条观测，覆盖小号铁钉、小号铁夹、普通钢筋和锈蚀钢筋四类样品；模型在对数衰减量上的 R²={m.get('r2_log')}，RMSE={m.get('rmse_log')}，换算到剩磁值后的 RMSE={m.get('rmse_m')} mT，MAPE={float(m.get('mape_m', '0')):.2%}。",
    )

    doc.add_heading("二、论文总体大纲", level=1)
    outline_rows = [
        ["摘要", "交代问题、模型体系、核心结果和判定流程", "图1、图6、图8；关键指标 R²/RMSE/MAPE"],
        ["1 问题重述", "从雷击后剩磁衰减、环境影响、阈值修正和现场判定四个任务重述题意", "题面、附件1、附件2"],
        ["2 数据说明与预处理", "说明样品类型、观测天数、天气变量、初值锚定和缺失异常处理", "图1、图2、表1"],
        ["3 模型假设与符号说明", "统一 M_sj(t)、M_sj(0)、R_s(t)、Y_s(t)、Theta_s(t) 等符号", "符号表"],
        ["4 问题一：动态衰减模型", "建立环境驱动指数衰减模型并预测第23-29天剩磁", "图1、图2、图3、预测表"],
        ["5 问题二：影响因素分析", "用主效应、交互项和消融分析解释时间、温湿度、锈蚀差异", "图4、图5、问题2结果表"],
        ["6 问题三：动态阈值修正", "把国标静态阈值映射为随检测延迟和环境变化衰减的动态阈值", "图6、阈值关键点表"],
        ["7 问题四：三区域判定流程", "基于反推初值区间和辅助概率输出支持、灰区、不支持", "图7、图8、问题4判定表"],
        ["8 检验、敏感性与稳健性", "汇总误差、单调性、消融、Bootstrap 区间宽度", "模型参数与检验、问题3单调性、问题4统计"],
        ["9 模型评价与推广", "指出可解释性、可操作性和局限；给出现场推广建议", "文字讨论"],
        ["10 结论", "凝练四问答案和工程化判定建议", "无"],
    ]
    add_table(doc, ["论文板块", "写作任务", "对应图表/数据"], outline_rows, [3.4, 7.1, 5.2])

    doc.add_heading("三、对应正文草稿", level=1)

    doc.add_heading("摘要（草稿）", level=2)
    add_paragraph(
        doc,
        "针对剩磁法雷击判定中剩磁随时间和环境衰减、固定阈值难以适应检测延迟的问题，本文以附件模拟实验数据和逐日天气数据为基础，建立了环境驱动的指数衰减模型。模型以样品第0天剩磁作为个体初值锚点，将相对衰减量表示为时间、累计温度、累计湿度及温湿耦合项的函数，并引入样品类型主效应和类型-时间交互项刻画不同材料的衰减差异。模型在对数衰减量上的拟合优度达到 R²=0.999409，剩磁尺度 RMSE 为 0.050741 mT，说明该模型能够较好重建四类样品的衰减轨迹。",
    )
    add_paragraph(
        doc,
        "在此基础上，本文进一步分析了关键影响因素，发现时间效应、锈蚀钢筋效应、累计湿度效应和温湿交互效应均显著；随后将国标静态阈值修正为动态阈值 Theta_s(t)=Theta_s(0)R_s(t)，得到四类样品 1-90 天的动态阈值序列。最后，本文构建了基于 Bootstrap 反推初值区间的三区域判定流程，将现场检测结果划分为支持曾遭雷击、灰区/证据不足和不支持曾遭雷击三类。结果表明，该流程能够同时输出判定结论、辅助概率和高风险提示，为雷击调查中的延迟检测提供了可复核、可解释的决策支持工具。",
    )
    add_paragraph(doc, "关键词：剩磁法；指数衰减；环境驱动；动态阈值；Bootstrap；三区域判定")

    doc.add_heading("1 问题重述（草稿）", level=2)
    add_paragraph(
        doc,
        "雷击会使铁磁性构件产生一定剩磁，但剩磁并非静态量，而会随放置时间、温度、湿度、材料状态及锈蚀程度变化而衰减。若直接采用固定阈值判断构件是否曾遭雷击，可能忽略检测滞后带来的系统性偏差。题目要求基于模拟实验数据和天气数据，建立剩磁衰减预测模型，分析影响因素，修正判定阈值，并形成可用于现场的雷击判定流程。",
    )
    add_paragraph(
        doc,
        "本文将四个问题统一到同一条建模链路中：首先根据样品初值和后续测量值建立动态衰减模型；其次对样品类型、时间、温湿度及交互效应进行显著性分析；再次将静态阈值转化为随时间衰减的动态阈值；最后利用观测剩磁反推初始剩磁区间，并据此构造三区域判定规则。",
    )

    doc.add_heading("2 数据说明与预处理（草稿）", level=2)
    add_paragraph(
        doc,
        "附件1包含四类样品的模拟实验观测，共1380条记录。小号铁钉和小号铁夹各包含460条记录，普通钢筋和锈蚀钢筋各包含230条记录；观测天数覆盖第0天至第90天。附件2给出对应时期逐日天气、气温和相对湿度，为刻画环境驱动衰减提供输入变量。",
    )
    add_paragraph(
        doc,
        "预处理时，以“样品类型+编号”作为个体标识，将第0天剩磁 M_sj(0) 作为个体初始剩磁，并构造相对衰减量 Y_sj(t)=ln[M_sj(0)/M_sj(t)]。这种处理可以剥离不同样品初始磁化水平的差异，使模型更专注于衰减规律本身。天气变量按测量天数匹配，并构造累计温度偏差、累计湿度偏差及累计温湿交互项，以反映环境的累积作用。",
    )

    doc.add_heading("3 模型假设与符号说明（草稿）", level=2)
    assumptions = [
        "同一类型样品共享主要衰减机制，个体差异主要体现在初始剩磁 M_sj(0) 上。",
        "剩磁衰减满足指数衰减框架，环境变量通过影响对数衰减量 Y_s(t) 改变保留率 R_s(t)。",
        "温度、湿度对剩磁衰减具有累积作用，且高温高湿可能产生耦合影响。",
        "静态阈值在检测延迟后应按同类型保留率进行动态修正。",
        "反推初值的不确定性可由样品簇 Bootstrap 近似刻画。"
    ]
    for item in assumptions:
        add_paragraph(doc, item)
    symbol_rows = [
        ["M_sj(t)", "类型 s 的第 j 个样品在第 t 天的剩磁"],
        ["M_sj(0)", "该样品第0天初始剩磁"],
        ["Y_s(t)", "类型 s 在第 t 天的对数衰减量"],
        ["R_s(t)", "剩磁保留率，R_s(t)=exp[-Y_s(t)]"],
        ["Theta_s(t)", "类型 s 在第 t 天的动态判定阈值"],
        ["CT(t), CH(t), CTH(t)", "累计温度、累计湿度和累计温湿交互变量"],
    ]
    add_table(doc, ["符号", "含义"], symbol_rows, [4.0, 11.5])

    doc.add_heading("4 问题一：动态衰减模型与预测（草稿）", level=2)
    add_paragraph(
        doc,
        "设 M_sj(t)=M_sj(0)exp[-Y_s(t)]。以普通钢筋作为基准类型，将样品类型虚拟变量、标准化时间、时间二次项、累计温度、累计湿度、累计温湿交互项以及类型-时间交互项纳入线性模型。该设计既保留指数衰减的物理可解释性，又允许不同材料在衰减速度上存在系统差异。",
    )
    add_paragraph(
        doc,
        f"模型结果显示，对数尺度 R²={m.get('r2_log')}，RMSE={m.get('rmse_log')}；换算回剩磁尺度后，RMSE={m.get('rmse_m')} mT，MAE={m.get('mae_m')} mT，MAPE={float(m.get('mape_m', '0')):.2%}。因此，该模型可作为后续预测、阈值修正和反推判定的公共基础模型。",
    )
    add_table(doc, list(pred.columns), pred.round(4).astype(str).values.tolist(), [2.0, 3.2, 3.2, 3.2, 3.2])

    doc.add_heading("5 问题二：关键影响因素分析（草稿）", level=2)
    add_paragraph(
        doc,
        "从主效应与交互项结果看，时间主效应和时间二次项均显著，说明剩磁衰减不仅随时间推进而增强，而且存在非线性趋势。锈蚀钢筋主效应显著为正，且锈蚀钢筋与时间的交互项系数为0.778232，表明锈蚀不仅改变初始衰减水平，还会放大长期检测延迟下的衰减风险。",
    )
    add_paragraph(
        doc,
        "环境变量方面，累计湿度主效应为正，累计温湿交互项也显著为正，说明湿度累积和高温高湿耦合会提高剩磁衰减程度。累计温度在去除时间趋势后的主效应为负，提示单独温度效应与时间趋势、湿度耦合之间存在方向差异，正文中应避免简单解释为“温度越高衰减越快”，而应强调环境变量需要在共同模型中解释。",
    )
    ablation_rows = p2_ablation.round(6).astype(str).values.tolist()
    add_table(doc, list(p2_ablation.columns), ablation_rows, [2.4, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0, 1.8])

    doc.add_heading("6 问题三：动态阈值修正（草稿）", level=2)
    add_paragraph(
        doc,
        "国标静态阈值适用于近似即时检测场景，但当检测发生在雷击后若干天，样品剩磁已经衰减，固定阈值会提高漏判风险。本文将静态阈值 Theta_s(0) 与保留率 R_s(t) 相乘，得到动态阈值 Theta_s(t)=Theta_s(0)R_s(t)。其中，小尺寸铁件采用 1.0 mT 作为初始阈值，钢筋类采用 1.5 mT 作为初始阈值。",
    )
    add_paragraph(
        doc,
        "动态阈值结果显示，四类样品阈值序列均单调不增；例如小号铁钉阈值由第1天0.875958 mT 下降到第90天0.000580 mT，锈蚀钢筋由第1天1.306901 mT 下降到第90天0.000227 mT。这说明检测延迟越长，若仍沿用静态阈值，越容易低估曾遭雷击的可能性。",
    )
    add_table(doc, list(p3_key.columns), p3_key.round(6).astype(str).values.tolist(), [1.4, 2.2, 1.8, 1.8, 2.5, 2.5, 2.5, 2.5, 2.6, 2.6, 2.6, 2.6])
    add_table(doc, list(p3_mono.columns), p3_mono.astype(str).values.tolist(), [2.8, 2.2, 2.0, 2.2, 2.2])

    doc.add_heading("7 问题四：三区域判定流程（草稿）", level=2)
    decision_counts = p4.iloc[:, 10].value_counts()
    high_counts = p4.iloc[:, 12].value_counts()
    add_paragraph(
        doc,
        f"问题四将现场检测值 M_obs 与模型保留率 R_s(t) 结合，反推初始剩磁 M0_hat=M_obs/R_s(t)，再用 Bootstrap 生成反推初值的95%区间。若区间下界高于静态阈值，则判为支持曾遭雷击；若区间上界低于阈值，则判为不支持；其余情况归入灰区/证据不足，并建议补充证据或复检。",
    )
    add_paragraph(
        doc,
        f"在当前模拟数据的1320条非零天观测中，{int(decision_counts.get('支持曾遭雷击', 0))}条被判为支持曾遭雷击，{int(decision_counts.get('灰区/证据不足', 0))}条进入灰区/证据不足；高风险提示中，{int(high_counts.get('是', 0))}条被标记为高风险。这一分布符合模拟实验以雷击后剩磁样本为主的背景，也表明长延迟、低保留率场景需要谨慎解释。",
    )

    doc.add_heading("8 模型检验、稳健性与图表组织（草稿）", level=2)
    add_paragraph(
        doc,
        "模型检验应从三个层面展开：第一，拟合误差层面报告对数尺度和剩磁尺度误差，说明模型对四类样品均有较小误差；第二，结构稳健性层面报告问题二消融分析，比较无交互、对称交互和门槛交互模型；第三，决策稳健性层面报告动态阈值单调性和反推初值区间宽度，说明判定流程在可解释性与不确定性控制之间取得平衡。",
    )
    add_paragraph(
        doc,
        "Academic Plotting 图表风格建议：正文图采用 figures2.1 的中文宋体版本，统一白底、浅灰网格、Nature/Science 风格配色；数值折线图优先展示95%置信区间，模型比较图展示误差棒或消融指标，流程图保持简洁，以便和论文正文的工程化判定流程对齐。",
    )
    figure_rows = [
        ["图1", "原始剩磁衰减曲线", "数据预处理与问题一开头", "展示四类样品剩磁随天数下降的总体趋势"],
        ["图2", "归一化保留率及95%置信区间", "问题一模型建立", "剥离初值差异，突出衰减规律"],
        ["图3", "初始剩磁与90天保留率关系", "问题一检验", "说明初值与长期保留率关系不应被简单混同"],
        ["图4", "主效应与交互效应系数图", "问题二因素分析", "展示关键变量方向、大小和显著性"],
        ["图5", "消融实验对比", "问题二稳健性", "比较交互项构造对误差的影响"],
        ["图6", "动态阈值曲线", "问题三核心结果", "展示固定阈值如何随检测延迟被修正"],
        ["图7", "反推初值区间宽度曲线", "问题四不确定性分析", "说明长延迟检测的不确定性变化"],
        ["图8", "现场判定流程图", "问题四方法总结", "把模型转化为可执行判定流程"],
    ]
    add_table(doc, ["编号", "图名", "建议位置", "写作作用"], figure_rows, [1.5, 4.2, 4.0, 6.0])

    doc.add_heading("9 模型评价与推广（草稿）", level=2)
    add_paragraph(
        doc,
        "本文模型的优点在于：一是以个体初值锚定消除了样品初始磁化差异，使衰减模型更稳定；二是把时间和环境累积效应纳入统一模型，具备较强可解释性；三是将阈值修正和三区域判定结合，能直接服务现场判定，并通过辅助概率和高风险提示表达不确定性。",
    )
    add_paragraph(
        doc,
        "模型的局限主要来自数据来源和适用范围。当前数据为模拟实验数据，天气变量只包含温度和湿度，尚未纳入风速、降雨强度、构件几何形状、磁测仪器误差等因素。若推广到真实现场，应进一步收集真实雷击案例和非雷击对照样本，并对动态阈值进行外部验证。",
    )

    doc.add_heading("10 结论（草稿）", level=2)
    add_paragraph(
        doc,
        "本文围绕剩磁法雷击判定预测问题，建立了环境驱动指数衰减模型，完成了剩磁预测、影响因素分析、动态阈值修正和三区域判定流程设计。结果表明，时间、锈蚀状态、湿度累积及温湿耦合是影响剩磁衰减的重要因素；动态阈值能够有效补偿检测延迟带来的阈值偏移；基于 Bootstrap 反推区间的三区域判定流程能够给出更稳健的现场决策建议。",
    )

    doc.add_heading("四、论文图表嵌入预览", level=1)
    fig_specs = [
        ("fig1_raw_decay_science.png", "图1 原始剩磁衰减曲线"),
        ("fig2_normalized_retention_science.png", "图2 归一化保留率及95%置信区间"),
        ("fig3_m0_vs_m90_ratio_science.png", "图3 初始剩磁与90天保留率关系"),
        ("fig4_problem2_coefficients_science.png", "图4 主效应与交互效应系数图"),
        ("fig5_problem2_ablation_science.png", "图5 消融实验对比"),
        ("fig6_dynamic_thresholds_science.png", "图6 动态阈值曲线"),
        ("fig7_interval_width_delay_science.png", "图7 反推初值区间宽度曲线"),
        ("fig8_flowchart_science.png", "图8 现场判定流程图"),
    ]
    for filename, caption in fig_specs:
        add_figure(doc, FIG / filename, caption)

    doc.add_heading("五、下一步写作建议", level=1)
    next_rows = [
        ["补齐公式编号", "将正文中的核心公式统一编号，保证符号表和正文一致。"],
        ["压缩表格", "正式论文中只保留关键表，完整 1-90 天阈值表放入附录。"],
        ["统一图题", "图题按“图x 主题”格式，正文首次引用每张图。"],
        ["加入参考文献", "补充剩磁法雷击判定、磁滞/剩磁衰减、Bootstrap 区间估计相关文献。"],
        ["按格式规范排版", "后续可基于比赛 Word 模板调整页眉、页脚、字体、摘要和正文格式。"],
    ]
    add_table(doc, ["任务", "说明"], next_rows, [4.0, 11.5])

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
