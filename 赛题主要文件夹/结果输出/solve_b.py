# -*- coding: utf-8 -*-
from __future__ import annotations

import csv
import math
import os
import random
import statistics as st
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from xml.sax.saxutils import escape


def find_base_dir() -> Path:
    env_path = os.environ.get("B_PATH")
    if env_path:
        return Path(env_path)
    return Path(__file__).resolve().parent.parent


def read_experiment(csv_path: Path):
    rows = []
    types_order = []
    with csv_path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.reader(f)
        next(reader)
        for r in reader:
            sample_type = r[0]
            if sample_type not in types_order:
                types_order.append(sample_type)
            rows.append(
                {
                    "type": sample_type,
                    "id": int(r[1]),
                    "day": int(r[2]),
                    "mag": float(r[3]),
                    "temp": float(r[4]),
                    "hum": float(r[5]),
                }
            )
    return rows, types_order


def col_to_idx(col: str) -> int:
    n = 0
    for ch in col:
        if ch.isalpha():
            n = n * 26 + ord(ch.upper()) - 64
    return n - 1


def read_xlsx_weather(xlsx_path: Path):
    with zipfile.ZipFile(xlsx_path) as zf:
        workbook = ET.fromstring(zf.read("xl/workbook.xml"))
        rels = ET.fromstring(zf.read("xl/_rels/workbook.xml.rels"))
        relmap = {rel.attrib["Id"]: rel.attrib["Target"] for rel in rels}
        sheets = workbook.find("{http://schemas.openxmlformats.org/spreadsheetml/2006/main}sheets")
        first_sheet = sheets[0]
        sheet_id = first_sheet.attrib[
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
        ]
        target = relmap[sheet_id].lstrip("/")
        if not target.startswith("xl/"):
            target = "xl/" + target

        shared = []
        if "xl/sharedStrings.xml" in zf.namelist():
            shared_xml = ET.fromstring(zf.read("xl/sharedStrings.xml"))
            ns_t = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}t"
            for si in shared_xml:
                shared.append("".join(t.text or "" for t in si.iter(ns_t)))

        rows = []
        worksheet = ET.fromstring(zf.read(target))
        ns_c = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}c"
        ns_v = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}v"
        ns_is = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}is"
        ns_t = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}t"
        for row in worksheet.iter("{http://schemas.openxmlformats.org/spreadsheetml/2006/main}row"):
            values = []
            for cell in row.iter(ns_c):
                ref = cell.attrib.get("r", "")
                idx = col_to_idx("".join(filter(str.isalpha, ref))) if ref else len(values)
                cell_type = cell.attrib.get("t")
                v = cell.find(ns_v)
                inline = cell.find(ns_is)
                value = ""
                if v is not None:
                    raw = v.text or ""
                    value = shared[int(raw)] if cell_type == "s" else raw
                elif inline is not None:
                    value = "".join(t.text or "" for t in inline.iter(ns_t))
                while len(values) <= idx:
                    values.append("")
                values[idx] = value
            if values and values[0].isdigit():
                rows.append(
                    {
                        "seq": int(values[0]),
                        "year": int(values[1]),
                        "month": int(values[2]),
                        "date": int(values[3]),
                        "weather": values[4],
                        "temp": float(values[5]),
                        "hum": float(values[6]),
                    }
                )
        return rows


def mean_sd(values):
    mean = sum(values) / len(values)
    sd = (sum((x - mean) ** 2 for x in values) / len(values)) ** 0.5
    return mean, sd


def invert_matrix(matrix):
    n = len(matrix)
    a = [row[:] + [1.0 if i == j else 0.0 for j in range(n)] for i, row in enumerate(matrix)]
    for col in range(n):
        pivot = max(range(col, n), key=lambda r: abs(a[r][col]))
        if abs(a[pivot][col]) < 1e-12:
            raise ValueError("Singular matrix in OLS fit")
        if pivot != col:
            a[col], a[pivot] = a[pivot], a[col]
        pv = a[col][col]
        for j in range(2 * n):
            a[col][j] /= pv
        for r in range(n):
            if r == col:
                continue
            factor = a[r][col]
            if factor:
                for j in range(2 * n):
                    a[r][j] -= factor * a[col][j]
    return [[a[i][j + n] for j in range(n)] for i in range(n)]


def fit_ols(x_rows, y):
    n = len(x_rows)
    p = len(x_rows[0])
    xtx = [[0.0] * p for _ in range(p)]
    xty = [0.0] * p
    for x, yi in zip(x_rows, y):
        for i in range(p):
            xty[i] += x[i] * yi
            for j in range(p):
                xtx[i][j] += x[i] * x[j]
    inv = invert_matrix(xtx)
    beta = [sum(inv[i][j] * xty[j] for j in range(p)) for i in range(p)]
    yhat = [sum(beta[j] * row[j] for j in range(p)) for row in x_rows]
    resid = [yi - yh for yi, yh in zip(y, yhat)]
    sse = sum(e * e for e in resid)
    ybar = sum(y) / n
    sst = sum((yi - ybar) ** 2 for yi in y)
    sigma2 = sse / (n - p)
    se = [(sigma2 * inv[i][i]) ** 0.5 for i in range(p)]
    return {
        "beta": beta,
        "se": se,
        "yhat": yhat,
        "resid": resid,
        "sse": sse,
        "rmse": (sse / n) ** 0.5,
        "r2": 1.0 - sse / sst,
        "sigma": sigma2**0.5,
    }


def fit_ridge(x_rows, y, lam: float = 1e-6):
    n = len(x_rows)
    p = len(x_rows[0])
    xtx = [[0.0] * p for _ in range(p)]
    xty = [0.0] * p
    for x, yi in zip(x_rows, y):
        for i in range(p):
            xty[i] += x[i] * yi
            for j in range(p):
                xtx[i][j] += x[i] * x[j]
    for i in range(p):
        xtx[i][i] += lam
    inv = invert_matrix(xtx)
    beta = [sum(inv[i][j] * xty[j] for j in range(p)) for i in range(p)]
    yhat = [sum(beta[j] * row[j] for j in range(p)) for row in x_rows]
    resid = [yi - yh for yi, yh in zip(y, yhat)]
    sse = sum(e * e for e in resid)
    return {
        "beta": beta,
        "yhat": yhat,
        "resid": resid,
        "sse": sse,
        "rmse": (sse / n) ** 0.5,
    }


def transpose_matrix(matrix):
    return [list(col) for col in zip(*matrix)]


def matrix_vector_product(matrix, vector):
    return [sum(a * b for a, b in zip(row, vector)) for row in matrix]


def cluster_robust_se(x_rows, resid, clusters):
    n = len(x_rows)
    p = len(x_rows[0])
    xtx = [[0.0] * p for _ in range(p)]
    for x in x_rows:
        for i in range(p):
            for j in range(p):
                xtx[i][j] += x[i] * x[j]
    xtx_inv = invert_matrix(xtx)
    cluster_map = {}
    for x, e, cluster in zip(x_rows, resid, clusters):
        bucket = cluster_map.setdefault(cluster, [0.0] * p)
        for i in range(p):
            bucket[i] += x[i] * e
    meat = [[0.0] * p for _ in range(p)]
    for vec in cluster_map.values():
        for i in range(p):
            for j in range(p):
                meat[i][j] += vec[i] * vec[j]
    cov = [[0.0] * p for _ in range(p)]
    for i in range(p):
        for j in range(p):
            cov[i][j] = sum(
                xtx_inv[i][k] * meat[k][l] * xtx_inv[l][j]
                for k in range(p)
                for l in range(p)
            )
    if n > p and len(cluster_map) > 1:
        g = len(cluster_map)
        scale = (g / (g - 1)) * ((n - 1) / (n - p))
        for i in range(p):
            for j in range(p):
                cov[i][j] *= scale
    se = [(cov[i][i] ** 0.5) if cov[i][i] >= 0 else float("nan") for i in range(p)]
    return se


def vif_values(x_rows, feature_indices=None):
    if feature_indices is None:
        feature_indices = list(range(1, len(x_rows[0])))
    selected = [[row[idx] for idx in feature_indices] for row in x_rows]
    vif = [float("nan")] * len(feature_indices)
    for j in range(len(feature_indices)):
        x_aux = [[1.0] + [row[k] for k in range(len(feature_indices)) if k != j] for row in selected]
        y_aux = [row[j] for row in selected]
        try:
            fit_aux = fit_ols(x_aux, y_aux)
            r2 = fit_aux["r2"]
            vif[j] = float("inf") if r2 >= 0.999999999 else 1.0 / max(1e-12, 1.0 - r2)
        except Exception:
            vif[j] = float("nan")
    return vif


def normal_p_value(t_value: float) -> float:
    return max(0.0, min(1.0, 2 * (1 - 0.5 * (1 + math.erf(abs(t_value) / math.sqrt(2))))))


def percentile(values, q: float) -> float:
    if not values:
        return float("nan")
    values = sorted(values)
    if len(values) == 1:
        return values[0]
    pos = (len(values) - 1) * q
    lo = int(math.floor(pos))
    hi = int(math.ceil(pos))
    if lo == hi:
        return values[lo]
    return values[lo] * (hi - pos) + values[hi] * (pos - lo)


def pearson_corr(xs, ys) -> float:
    if len(xs) != len(ys) or len(xs) < 2:
        return float("nan")
    mx = sum(xs) / len(xs)
    my = sum(ys) / len(ys)
    num = sum((x - mx) * (y - my) for x, y in zip(xs, ys))
    den_x = sum((x - mx) ** 2 for x in xs)
    den_y = sum((y - my) ** 2 for y in ys)
    den = (den_x * den_y) ** 0.5
    return num / den if den else float("nan")


def standardize(values):
    mean, sd = mean_sd(values)
    if abs(sd) < 1e-12:
        return [0.0 for _ in values]
    return [(value - mean) / sd for value in values]


def residualize(target, base_rows):
    fit = fit_ols(base_rows, target)
    return [value - fitted for value, fitted in zip(target, fit["yhat"])]


def residualize_and_standardize(target, base_rows):
    return standardize(residualize(target, base_rows))


def write_csv(path: Path, headers, rows):
    with open_output_text(path, encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        if headers:
            writer.writerow(headers)
        writer.writerows(rows)


def summarize_prediction_errors(observed_rows, yhat, m0, types_order):
    by_type = {sample_type: {"sq": [], "ab": [], "ape": []} for sample_type in types_order}
    all_sq, all_ab, all_ape = [], [], []
    for row, pred_log in zip(observed_rows, yhat):
        pred = m0[(row["type"], row["id"])] * math.exp(-pred_log)
        err = row["mag"] - pred
        bucket = by_type[row["type"]]
        bucket["sq"].append(err * err)
        bucket["ab"].append(abs(err))
        bucket["ape"].append(abs(err) / row["mag"])
        all_sq.append(err * err)
        all_ab.append(abs(err))
        all_ape.append(abs(err) / row["mag"])
    overall = {
        "rmse_m": (sum(all_sq) / len(all_sq)) ** 0.5 if all_sq else float("nan"),
        "mae_m": sum(all_ab) / len(all_ab) if all_ab else float("nan"),
        "mape_m": sum(all_ape) / len(all_ape) if all_ape else float("nan"),
    }
    by_type_summary = {}
    for sample_type, bucket in by_type.items():
        if not bucket["sq"]:
            by_type_summary[sample_type] = {"rmse": float("nan"), "mae": float("nan"), "mape": float("nan")}
            continue
        by_type_summary[sample_type] = {
            "rmse": (sum(bucket["sq"]) / len(bucket["sq"])) ** 0.5,
            "mae": sum(bucket["ab"]) / len(bucket["ab"]),
            "mape": sum(bucket["ape"]) / len(bucket["ape"]),
        }
    return overall, by_type_summary


def col_letter(index: int) -> str:
    letters = ""
    index += 1
    while index:
        index, rem = divmod(index - 1, 26)
        letters = chr(65 + rem) + letters
    return letters


def write_simple_xlsx(path: Path, headers, rows):
    def cell_xml(value, r):
        if isinstance(value, (int, float)) and not isinstance(value, bool):
            if isinstance(value, float):
                text = f"{value:.10g}"
            else:
                text = str(value)
            return f'<c r="{r}"><v>{text}</v></c>'
        text = escape(str(value))
        return f'<c r="{r}" t="inlineStr"><is><t>{text}</t></is></c>'

    sheet_rows = []
    all_rows = [headers] + rows
    for ridx, row in enumerate(all_rows, start=1):
        cells = []
        for cidx, value in enumerate(row):
            cells.append(cell_xml(value, f"{col_letter(cidx)}{ridx}"))
        sheet_rows.append(f'<row r="{ridx}">{"".join(cells)}</row>')
    worksheet = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        "<sheetData>"
        + "".join(sheet_rows)
        + "</sheetData></worksheet>"
    )
    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/xl/workbook.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
        '<Override PartName="/xl/worksheets/sheet1.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
        "</Types>"
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="xl/workbook.xml"/></Relationships>'
    )
    workbook = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<sheets><sheet name="Sheet1" sheetId="1" r:id="rId1"/></sheets></workbook>'
    )
    workbook_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
        'Target="worksheets/sheet1.xml"/></Relationships>'
    )
    with zipfile.ZipFile(resolve_output_path(path), "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("xl/workbook.xml", workbook)
        zf.writestr("xl/_rels/workbook.xml.rels", workbook_rels)
        zf.writestr("xl/worksheets/sheet1.xml", worksheet)


def resolve_output_path(path: Path) -> Path:
    candidates = [path]
    for idx in range(1, 100):
        candidates.append(path.with_name(f"{path.stem}_retry{idx}{path.suffix}"))
    for candidate in candidates:
        try:
            with candidate.open("a", encoding="utf-8", newline=""):
                return candidate
        except (PermissionError, OSError):
            continue
    raise PermissionError(f"Unable to open output path for writing: {path}")


def open_output_text(path: Path, encoding="utf-8-sig", newline=""):
    return resolve_output_path(path).open("w", encoding=encoding, newline=newline)


def main():
    base = find_base_dir()
    out = base / "B题_solution_outputs"
    out.mkdir(exist_ok=True)
    csv_path = next(p for p in base.iterdir() if p.suffix.lower() == ".csv")
    xlsx_path = next(p for p in base.iterdir() if p.suffix.lower() == ".xlsx")

    observations, types_order = read_experiment(csv_path)
    nail, clip, ordinary, rusty = types_order
    weather_rows = read_xlsx_weather(xlsx_path)
    weather = {r["seq"]: r for r in weather_rows}

    temps = [weather[d]["temp"] for d in range(1, 91)]
    hums = [weather[d]["hum"] for d in range(1, 91)]
    temp_mean = sum(temps) / len(temps)
    hum_mean = sum(hums) / len(hums)

    cum_temp, cum_hum, cum_inter = [], [], []
    cum_thr = []
    a = b = c = 0.0
    dthr = 0.0
    for d in range(1, 91):
        tc = weather[d]["temp"] - temp_mean
        hc = weather[d]["hum"] - hum_mean
        a += tc
        b += hc
        c += tc * hc
        dthr += max(weather[d]["temp"] - temp_mean, 0) * max(weather[d]["hum"] - hum_mean, 0)
        cum_temp.append(a)
        cum_hum.append(b)
        cum_inter.append(c)
        cum_thr.append(dthr)

    day_mean, day_sd = mean_sd(list(range(1, 91)))
    ct_mean, ct_sd = mean_sd(cum_temp)
    ch_mean, ch_sd = mean_sd(cum_hum)
    cth_mean, cth_sd = mean_sd(cum_inter)
    cthr_mean, cthr_sd = mean_sd(cum_thr)

    m0 = {}
    for row in observations:
        if row["day"] == 0:
            m0[(row["type"], row["id"])] = row["mag"]

    def features(sample_type: str, day: int):
        dz = (day - day_mean) / day_sd
        ct = (cum_temp[day - 1] - ct_mean) / ct_sd
        ch = (cum_hum[day - 1] - ch_mean) / ch_sd
        cth = (cum_inter[day - 1] - cth_mean) / cth_sd
        n = 1.0 if sample_type == nail else 0.0
        cl = 1.0 if sample_type == clip else 0.0
        r = 1.0 if sample_type == rusty else 0.0
        return [1.0, n, cl, r, dz, dz * dz, ct, ch, cth, r * dz, n * dz, cl * dz]

    def features_ablation(sample_type: str, day: int, interaction_mode: str):
        dz = (day - day_mean) / day_sd
        ct = (cum_temp[day - 1] - ct_mean) / ct_sd
        ch = (cum_hum[day - 1] - ch_mean) / ch_sd
        cth = (cum_inter[day - 1] - cth_mean) / cth_sd
        cthr = (cum_thr[day - 1] - cthr_mean) / cthr_sd
        n = 1.0 if sample_type == nail else 0.0
        cl = 1.0 if sample_type == clip else 0.0
        r = 1.0 if sample_type == rusty else 0.0
        base = [1.0, n, cl, r, dz, dz * dz, ct, ch, r * dz, n * dz, cl * dz]
        if interaction_mode == "none":
            return base
        if interaction_mode == "symmetric":
            return base + [cth]
        if interaction_mode == "threshold":
            return base + [cthr]
        raise ValueError(interaction_mode)

    def simple_features(day: int):
        dz = (day - day_mean) / day_sd
        ct = (cum_temp[day - 1] - ct_mean) / ct_sd
        ch = (cum_hum[day - 1] - ch_mean) / ch_sd
        cth = (cum_inter[day - 1] - cth_mean) / cth_sd
        return [1.0, dz, dz * dz, ct, ch, cth]

    def sample_flags(sample_type: str):
        return (
            1.0 if sample_type == nail else 0.0,
            1.0 if sample_type == clip else 0.0,
            1.0 if sample_type == rusty else 0.0,
        )

    theta0_by_type = {
        nail: 1.0,
        clip: 1.0,
        ordinary: 1.5,
        rusty: 1.5,
    }

    problem2_days = list(range(1, 91))
    problem2_day_z = [(day - day_mean) / day_sd for day in problem2_days]
    problem2_day_z2 = [value * value for value in problem2_day_z]
    problem2_projection_base = [
        [1.0, dz, dz2] for dz, dz2 in zip(problem2_day_z, problem2_day_z2)
    ]
    problem2_cum_temp_z = residualize_and_standardize(cum_temp, problem2_projection_base)
    problem2_cum_hum_z = residualize_and_standardize(cum_hum, problem2_projection_base)
    problem2_cum_inter_sym_z = residualize_and_standardize(cum_inter, problem2_projection_base)
    problem2_cum_inter_thr_z = residualize_and_standardize(cum_thr, problem2_projection_base)

    def features_problem2(sample_type: str, day: int, interaction_mode: str = "symmetric"):
        nail_flag, clip_flag, rusty_flag = sample_flags(sample_type)
        dz = problem2_day_z[day - 1]
        dz2 = problem2_day_z2[day - 1]
        ct = problem2_cum_temp_z[day - 1]
        ch = problem2_cum_hum_z[day - 1]
        cth_sym = problem2_cum_inter_sym_z[day - 1]
        cth_thr = problem2_cum_inter_thr_z[day - 1]
        base = [
            1.0,
            nail_flag,
            clip_flag,
            rusty_flag,
            dz,
            dz2,
            ct,
            ch,
            rusty_flag * dz,
            nail_flag * dz,
            clip_flag * dz,
        ]
        if interaction_mode == "none":
            return base
        if interaction_mode == "symmetric":
            return base + [cth_sym]
        if interaction_mode == "threshold":
            return base + [cth_thr]
        raise ValueError(interaction_mode)

    def features_problem2_reduced(sample_type: str, day: int):
        nail_flag, clip_flag, rusty_flag = sample_flags(sample_type)
        dz = problem2_day_z[day - 1]
        ct = problem2_cum_temp_z[day - 1]
        ch = problem2_cum_hum_z[day - 1]
        return [1.0, nail_flag, clip_flag, rusty_flag, dz, ct, ch, rusty_flag * dz]

    labels = [
        "截距",
        "小号铁钉",
        "小号铁夹",
        "锈蚀钢筋",
        "时间_z",
        "时间_z^2",
        "累计温度_z",
        "累计湿度_z",
        "累计温湿交互_z",
        "锈蚀钢筋×时间_z",
        "小号铁钉×时间_z",
        "小号铁夹×时间_z",
    ]

    x_rows, y = [], []
    for row in observations:
        if row["day"] == 0:
            continue
        x_rows.append(features(row["type"], row["day"]))
        y.append(math.log(m0[(row["type"], row["id"])] / row["mag"]))
    model = fit_ols(x_rows, y)
    ridge_model = fit_ridge(x_rows, y, lam=0.1)

    def predict_log_decay(sample_type: str, day: int) -> float:
        x = features(sample_type, day)
        return sum(bi * xi for bi, xi in zip(model["beta"], x))

    def survival(sample_type: str, day: int) -> float:
        return math.exp(-predict_log_decay(sample_type, day))

    initial_stats = {}
    for sample_type in types_order:
        vals = [value for (typ, _), value in m0.items() if typ == sample_type]
        initial_stats[sample_type] = {
            "mean": st.mean(vals),
            "sd": st.pstdev(vals),
            "min": min(vals),
            "max": max(vals),
            "n": len(vals),
        }

    survival_cache_raw = {
        (sample_type, day): survival(sample_type, day)
        for sample_type in types_order
        for day in range(1, 91)
    }

    survival_cache = {}
    for sample_type in types_order:
        running_min = float("inf")
        for day in range(1, 91):
            running_min = min(running_min, survival_cache_raw[(sample_type, day)])
            survival_cache[(sample_type, day)] = running_min

    def theta_obs(sample_type: str, day: int, use_mean: bool = False) -> float:
        base = initial_stats[sample_type]["mean"] if use_mean else theta0_by_type[sample_type]
        return round(base * survival_cache[(sample_type, day)], 6)

    def theta_static(sample_type: str) -> float:
        return theta0_by_type[sample_type]

    threshold_profile_by_type = {}
    for sample_type in types_order:
        series = []
        running_min = float("inf")
        for day in range(1, 91):
            running_min = min(running_min, theta_obs(sample_type, day))
            series.append(running_min)
        threshold_profile_by_type[sample_type] = series

    metrics_by_type = {}
    all_sq, all_abs, all_ape = [], [], []
    for sample_type in types_order:
        sq, ab, ape = [], [], []
        for row in observations:
            if row["day"] == 0 or row["type"] != sample_type:
                continue
            pred = m0[(row["type"], row["id"])] * survival(sample_type, row["day"])
            err = row["mag"] - pred
            sq.append(err * err)
            ab.append(abs(err))
            ape.append(abs(err) / row["mag"])
        all_sq.extend(sq)
        all_abs.extend(ab)
        all_ape.extend(ape)
        metrics_by_type[sample_type] = {
            "rmse": (sum(sq) / len(sq)) ** 0.5,
            "mae": sum(ab) / len(ab),
            "mape": sum(ape) / len(ape),
        }
    metrics = {
        "r2_log": model["r2"],
        "rmse_log": model["rmse"],
        "sigma_log": model["sigma"],
        "rmse_m": (sum(all_sq) / len(all_sq)) ** 0.5,
        "mae_m": sum(all_abs) / len(all_abs),
        "mape_m": sum(all_ape) / len(all_ape),
    }

    continuous_diag_idx = [4, 5, 6, 7, 8]
    problem2_collinearity_path = out / "问题2_共线性诊断.csv"
    with open_output_text(problem2_collinearity_path, encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["变量A", "变量B", "相关系数"])
        for i in range(len(continuous_diag_idx)):
            for j in range(i + 1, len(continuous_diag_idx)):
                ia = continuous_diag_idx[i]
                ib = continuous_diag_idx[j]
                xs = [row[ia] for row in x_rows]
                ys = [row[ib] for row in x_rows]
                writer.writerow([labels[ia], labels[ib], round(pearson_corr(xs, ys), 6)])

    problem2_ridge_path = out / "问题2_岭回归稳健性.csv"
    with open_output_text(problem2_ridge_path, encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["变量", "OLS系数", "Ridge系数", "符号一致"])
        for label, beta_ols, beta_ridge in zip(labels, model["beta"], ridge_model["beta"]):
            same_sign = "是" if beta_ols == 0 or beta_ridge == 0 or (beta_ols > 0) == (beta_ridge > 0) else "否"
            writer.writerow([label, round(beta_ols, 6), round(beta_ridge, 6), same_sign])
        writer.writerow([])
        writer.writerow(["指标", "OLS", "Ridge"])
        writer.writerow(["RMSE_log", round(model["rmse"], 6), round(ridge_model["rmse"], 6)])
        writer.writerow(["R2_log", round(model["r2"], 6), "NA"])

    observed_rows_day = [row for row in observations if row["day"] > 0]
    problem2_labels = [
        "截距",
        "小号铁钉",
        "小号铁夹",
        "锈蚀钢筋",
        "时间_z",
        "时间_z^2",
        "累计温度_z",
        "累计湿度_z",
        "锈蚀钢筋×时间_z",
        "小号铁钉×时间_z",
        "小号铁夹×时间_z",
        "累计温湿交互_z",
    ]
    x_rows_problem2 = [features_problem2(row["type"], row["day"], "symmetric") for row in observed_rows_day]
    y_problem2 = [math.log(m0[(row["type"], row["id"])] / row["mag"]) for row in observed_rows_day]
    problem2_model = fit_ols(x_rows_problem2, y_problem2)
    problem2_ridge_model = fit_ridge(x_rows_problem2, y_problem2, lam=0.1)
    x_rows_problem2_reduced = [features_problem2_reduced(row["type"], row["day"]) for row in observed_rows_day]
    problem2_reduced_model = fit_ols(x_rows_problem2_reduced, y_problem2)
    problem2_clusters = [(row["type"], row["id"]) for row in observed_rows_day]
    problem2_cluster_se = cluster_robust_se(x_rows_problem2, problem2_model["resid"], problem2_clusters)
    problem2_vif_indices = [4, 5, 6, 7, 11]
    problem2_vif_labels = [problem2_labels[i] for i in problem2_vif_indices]
    problem2_vif = vif_values(x_rows_problem2, problem2_vif_indices)

    def cross_validate_features(feature_builder):
        cv_sq, cv_ab, cv_ape = [], [], []
        for sample_type in types_order:
            sample_ids = sorted({row["id"] for row in observations if row["type"] == sample_type and row["day"] == 0})
            for sample_id in sample_ids:
                train_rows = [row for row in observed_rows_day if not (row["type"] == sample_type and row["id"] == sample_id)]
                x_train = [feature_builder(row["type"], row["day"]) for row in train_rows]
                y_train = [math.log(m0[(row["type"], row["id"])] / row["mag"]) for row in train_rows]
                fit_cv = fit_ols(x_train, y_train)
                held_rows = [row for row in observed_rows_day if row["type"] == sample_type and row["id"] == sample_id]
                for row in held_rows:
                    pred_log = sum(beta * value for beta, value in zip(fit_cv["beta"], feature_builder(row["type"], row["day"])))
                    pred = m0[(row["type"], row["id"])] * math.exp(-pred_log)
                    err = row["mag"] - pred
                    cv_sq.append(err * err)
                    cv_ab.append(abs(err))
                    cv_ape.append(abs(err) / row["mag"])
        return {
            "rmse": (sum(cv_sq) / len(cv_sq)) ** 0.5 if cv_sq else float("nan"),
            "mae": sum(cv_ab) / len(cv_ab) if cv_ab else float("nan"),
            "mape": sum(cv_ape) / len(cv_ape) if cv_ape else float("nan"),
        }

    problem2_shrink_path = out / "问题2_收缩模型对比.csv"
    problem2_shrink_rows = []
    full_cv = cross_validate_features(lambda s, d: features_problem2(s, d, "symmetric"))
    reduced_cv = cross_validate_features(features_problem2_reduced)
    full_aic = len(y_problem2) * math.log(problem2_model["sse"] / len(y_problem2)) + 2 * len(problem2_model["beta"])
    reduced_aic = len(y_problem2) * math.log(problem2_reduced_model["sse"] / len(y_problem2)) + 2 * len(problem2_reduced_model["beta"])
    full_bic = len(y_problem2) * math.log(problem2_model["sse"] / len(y_problem2)) + len(problem2_model["beta"]) * math.log(len(y_problem2))
    reduced_bic = len(y_problem2) * math.log(problem2_reduced_model["sse"] / len(y_problem2)) + len(problem2_reduced_model["beta"]) * math.log(len(y_problem2))
    problem2_shrink_rows.append(["完整模型", round(problem2_model["rmse"], 6), round(full_cv["rmse"], 6), round(full_aic, 6), round(full_bic, 6), len(problem2_model["beta"])])
    problem2_shrink_rows.append(["收缩模型", round(problem2_reduced_model["rmse"], 6), round(reduced_cv["rmse"], 6), round(reduced_aic, 6), round(reduced_bic, 6), len(problem2_reduced_model["beta"])])
    problem2_shrink_rows.append(["CV_RMSE差值(完整-收缩)", round(full_cv["rmse"] - reduced_cv["rmse"], 6), "", "", "", ""])
    problem2_shrink_rows.append(["BIC差值(完整-收缩)", "", "", round(full_bic - reduced_bic, 6), "", ""])
    write_csv(problem2_shrink_path, ["模型", "RMSE_log", "CV_RMSE_log", "AIC", "BIC", "参数个数"], problem2_shrink_rows)

    problem2_collinearity_path = out / "问题2_共线性诊断.csv"
    problem2_collinearity_rows = []
    problem2_continuous_idx = [4, 5, 6, 7, 11]
    for i in range(len(problem2_continuous_idx)):
        for j in range(i + 1, len(problem2_continuous_idx)):
            ia = problem2_continuous_idx[i]
            ib = problem2_continuous_idx[j]
            xs = [row[ia] for row in x_rows_problem2]
            ys = [row[ib] for row in x_rows_problem2]
            problem2_collinearity_rows.append(
                [problem2_labels[ia], problem2_labels[ib], round(pearson_corr(xs, ys), 6)]
            )
    write_csv(problem2_collinearity_path, ["变量A", "变量B", "相关系数"], problem2_collinearity_rows)

    problem2_ridge_path = out / "问题2_岭回归稳健性.csv"
    problem2_ridge_rows = []
    for label, beta_ols, beta_ridge in zip(problem2_labels, problem2_model["beta"], problem2_ridge_model["beta"]):
        same_sign = "是" if beta_ols == 0 or beta_ridge == 0 or (beta_ols > 0) == (beta_ridge > 0) else "否"
        problem2_ridge_rows.append([label, round(beta_ols, 6), round(beta_ridge, 6), same_sign])
    problem2_ridge_rows.append([])
    problem2_ridge_rows.append(["指标", "OLS", "Ridge"])
    problem2_ridge_rows.append(["RMSE_log", round(problem2_model["rmse"], 6), round(problem2_ridge_model["rmse"], 6)])
    problem2_ridge_rows.append(["R2_log", round(problem2_model["r2"], 6), "NA"])
    write_csv(problem2_ridge_path, ["变量", "OLS系数", "Ridge系数", "符号一致"], problem2_ridge_rows)

    problem2_vif_path = out / "问题2_VIF诊断.csv"
    problem2_vif_rows = []
    for label, vif in zip(problem2_vif_labels, problem2_vif):
        problem2_vif_rows.append([label, round(vif, 6) if not math.isnan(vif) else "NA"])
    write_csv(problem2_vif_path, ["变量", "VIF"], problem2_vif_rows)

    # Problem 1 foundation tests: initial-value independence and shared-kernel comparison.
    def sample_rows(sample_type: str, sample_id: int):
        return [row for row in observations if row["type"] == sample_type and row["id"] == sample_id]

    def fit_type_shared_model(sample_type: str, excluded_ids=None):
        excluded_ids = set(excluded_ids or [])
        x_rows_t, y_t = [], []
        for row in observations:
            if row["day"] == 0 or row["type"] != sample_type or row["id"] in excluded_ids:
                continue
            m0_val = m0[(sample_type, row["id"])]
            x_rows_t.append(simple_features(row["day"]))
            y_t.append(math.log(m0_val / row["mag"]))
        return fit_ols(x_rows_t, y_t)

    problem1_rows = []

    for sample_type in types_order:
        type_corr_x, type_corr_y = [], []
        type_shared_x, type_shared_y = [], []
        type_ind_sse = type_ind_n = type_ind_p = 0.0
        sample_ids = sorted({row["id"] for row in observations if row["type"] == sample_type and row["day"] == 0})
        for sample_id in sample_ids:
            rows_i = sample_rows(sample_type, sample_id)
            m0_val = m0[(sample_type, sample_id)]
            m90_val = next(row["mag"] for row in rows_i if row["day"] == 90)
            type_corr_x.append(m0_val)
            type_corr_y.append(m90_val / m0_val)

            x_i, y_i = [], []
            for row in rows_i:
                if row["day"] == 0:
                    continue
                x_i.append(simple_features(row["day"]))
                y_i.append(math.log(m0_val / row["mag"]))
                type_shared_x.append(simple_features(row["day"]))
                type_shared_y.append(math.log(m0_val / row["mag"]))

            fit_i = fit_ols(x_i, y_i)
            type_ind_sse += fit_i["sse"]
            type_ind_n += len(y_i)
            type_ind_p += len(fit_i["beta"])

        shared_fit = fit_type_shared_model(sample_type)
        shared_n = len(type_shared_y)
        shared_p = len(shared_fit["beta"])
        shared_aic = shared_n * math.log(shared_fit["sse"] / shared_n) + 2 * shared_p
        shared_bic = shared_n * math.log(shared_fit["sse"] / shared_n) + shared_p * math.log(shared_n)
        ind_aic = type_ind_n * math.log(type_ind_sse / type_ind_n) + 2 * type_ind_p
        ind_bic = type_ind_n * math.log(type_ind_sse / type_ind_n) + type_ind_p * math.log(type_ind_n)
        problem1_rows.append(
            [
                sample_type,
                len(sample_ids),
                round(pearson_corr(type_corr_x, type_corr_y), 6),
                round(shared_fit["rmse"], 6),
                round((type_ind_sse / type_ind_n) ** 0.5, 6),
                round(shared_aic, 6),
                round(ind_aic, 6),
                round(shared_bic, 6),
                round(ind_bic, 6),
            ]
        )

    problem1_path = out / "问题1_地基检验结果.csv"
    with open_output_text(problem1_path, encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["样品类型", "样本数", "corr(M0,M90/M0)", "共享模型RMSE", "独立模型RMSE", "共享模型AIC", "独立模型AIC", "共享模型BIC", "独立模型BIC"])
        writer.writerows(problem1_rows)

    problem1_cv_rows = []
    overall_cv_sq, overall_cv_ab, overall_cv_ape = [], [], []
    for sample_type in types_order:
        sample_ids = sorted({row["id"] for row in observations if row["type"] == sample_type and row["day"] == 0})
        cv_sq, cv_ab, cv_ape = [], [], []
        for sample_id in sample_ids:
            fit_cv = fit_type_shared_model(sample_type, excluded_ids={sample_id})
            beta_cv = fit_cv["beta"]
            held_rows = sample_rows(sample_type, sample_id)
            m0_val = m0[(sample_type, sample_id)]
            for row in held_rows:
                if row["day"] == 0:
                    continue
                yhat = sum(bi * xi for bi, xi in zip(beta_cv, simple_features(row["day"])))
                pred = m0_val * math.exp(-yhat)
                err = row["mag"] - pred
                cv_sq.append(err * err)
                cv_ab.append(abs(err))
                cv_ape.append(abs(err) / row["mag"])
        overall_cv_sq.extend(cv_sq)
        overall_cv_ab.extend(cv_ab)
        overall_cv_ape.extend(cv_ape)
        problem1_cv_rows.append(
            [
                sample_type,
                len(sample_ids),
                round((sum(cv_sq) / len(cv_sq)) ** 0.5, 6),
                round(sum(cv_ab) / len(cv_ab), 6),
                round(sum(cv_ape) / len(cv_ape), 6),
            ]
        )
    problem1_cv_rows.append(
        [
            "总体",
            len({(row["type"], row["id"]) for row in observations if row["day"] == 0}),
            round((sum(overall_cv_sq) / len(overall_cv_sq)) ** 0.5, 6),
            round(sum(overall_cv_ab) / len(overall_cv_ab), 6),
            round(sum(overall_cv_ape) / len(overall_cv_ape), 6),
        ]
    )

    problem1_cv_path = out / "问题1_LOSO交叉验证.csv"
    with open_output_text(problem1_cv_path, encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["样品类型", "样本数", "CV_RMSE(mT)", "CV_MAE(mT)", "CV_MAPE"])
        writer.writerows(problem1_cv_rows)

    # Problem 2 ablation analysis.
    variant_specs = [
        ("无交互", "none", ["截距", "小号铁钉", "小号铁夹", "锈蚀钢筋", "时间_z", "时间_z^2", "累计温度_z", "累计湿度_z", "锈蚀钢筋×时间_z", "小号铁钉×时间_z", "小号铁夹×时间_z"]),
        ("对称交互", "symmetric", ["截距", "小号铁钉", "小号铁夹", "锈蚀钢筋", "时间_z", "时间_z^2", "累计温度_z", "累计湿度_z", "锈蚀钢筋×时间_z", "小号铁钉×时间_z", "小号铁夹×时间_z", "累计温湿交互_z"]),
        ("门槛交互", "threshold", ["截距", "小号铁钉", "小号铁夹", "锈蚀钢筋", "时间_z", "时间_z^2", "累计温度_z", "累计湿度_z", "锈蚀钢筋×时间_z", "小号铁钉×时间_z", "小号铁夹×时间_z", "门槛型温湿交互_z"]),
    ]

    problem2_summary_rows = []
    problem2_detail_rows = []
    problem2_cv_rows = []
    for variant_name, mode, variant_labels in variant_specs:
        x_rows_v, y_v = [], []
        for row in observations:
            if row["day"] == 0:
                continue
            x_rows_v.append(features_ablation(row["type"], row["day"], mode))
            y_v.append(math.log(m0[(row["type"], row["id"])] / row["mag"]))
        fit_v = fit_ols(x_rows_v, y_v)
        yhat_v = fit_v["yhat"]
        pred_by_type = {sample_type: {"sq": [], "ab": [], "ape": []} for sample_type in types_order}
        for row, yhat in zip([r for r in observations if r["day"] > 0], yhat_v):
            pred = m0[(row["type"], row["id"])] * math.exp(-yhat)
            err = row["mag"] - pred
            bucket = pred_by_type[row["type"]]
            bucket["sq"].append(err * err)
            bucket["ab"].append(abs(err))
            bucket["ape"].append(abs(err) / row["mag"])
        key_beta = None
        key_p = None
        if mode == "symmetric":
            key_beta = fit_v["beta"][-1]
            key_p = normal_p_value(key_beta / fit_v["se"][-1])
        elif mode == "threshold":
            key_beta = fit_v["beta"][-1]
            key_p = normal_p_value(key_beta / fit_v["se"][-1])
        problem2_summary_rows.append(
            [
                variant_name,
                round(fit_v["r2"], 6),
                round(fit_v["rmse"], 6),
                round((sum(sum(v["sq"]) for v in pred_by_type.values()) / sum(len(v["sq"]) for v in pred_by_type.values())) ** 0.5, 6),
                round(sum(sum(v["ab"]) for v in pred_by_type.values()) / sum(len(v["ab"]) for v in pred_by_type.values()), 6),
                round(sum(sum(v["ape"]) for v in pred_by_type.values()) / sum(len(v["ape"]) for v in pred_by_type.values()), 6),
                round(key_beta, 6) if key_beta is not None else "",
                round(key_p, 6) if key_p is not None else "",
            ]
        )
        for sample_type in types_order:
            bucket = pred_by_type[sample_type]
            problem2_detail_rows.append(
                [
                    variant_name,
                    sample_type,
                    round((sum(bucket["sq"]) / len(bucket["sq"])) ** 0.5, 6),
                    round(sum(bucket["ab"]) / len(bucket["ab"]), 6),
                    round(sum(bucket["ape"]) / len(bucket["ape"]), 6),
                ]
            )

        cv_sq, cv_ab, cv_ape = [], [], []
        for sample_type in types_order:
            sample_ids = sorted({row["id"] for row in observations if row["type"] == sample_type and row["day"] == 0})
            for sample_id in sample_ids:
                train_rows = [
                    row for row in observations
                    if row["day"] > 0 and not (row["type"] == sample_type and row["id"] == sample_id)
                ]
                if len(train_rows) <= len(variant_labels):
                    continue
                x_train, y_train = [], []
                for row in train_rows:
                    x_train.append(features_ablation(row["type"], row["day"], mode))
                    y_train.append(math.log(m0[(row["type"], row["id"])] / row["mag"]))
                try:
                    cv_fit = fit_ols(x_train, y_train)
                except ValueError:
                    continue
                held_rows = [row for row in observations if row["type"] == sample_type and row["id"] == sample_id and row["day"] > 0]
                for row in held_rows:
                    yhat = sum(bi * xi for bi, xi in zip(cv_fit["beta"], features_ablation(row["type"], row["day"], mode)))
                    pred = m0[(row["type"], row["id"])] * math.exp(-yhat)
                    err = row["mag"] - pred
                    cv_sq.append(err * err)
                    cv_ab.append(abs(err))
                    cv_ape.append(abs(err) / row["mag"])

        problem2_cv_rows.append(
            [
                variant_name,
                round((sum(cv_sq) / len(cv_sq)) ** 0.5, 6) if cv_sq else "",
                round(sum(cv_ab) / len(cv_ab), 6) if cv_ab else "",
                round(sum(cv_ape) / len(cv_ape), 6) if cv_ape else "",
            ]
        )

    problem2_summary_path = out / "问题2_消融分析.csv"
    with open_output_text(problem2_summary_path, encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["模型版本", "R2_log", "RMSE_log", "RMSE_m", "MAE_m", "MAPE_m", "末项系数", "末项p值"])
        writer.writerows(problem2_summary_rows)

    problem2_detail_path = out / "问题2_类型误差对比.csv"
    with open_output_text(problem2_detail_path, encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["模型版本", "样品类型", "RMSE(mT)", "MAE(mT)", "MAPE"])
        writer.writerows(problem2_detail_rows)

    problem2_cv_path = out / "问题2_交叉验证.csv"
    with open_output_text(problem2_cv_path, encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["模型版本", "CV_RMSE(mT)", "CV_MAE(mT)", "CV_MAPE"])
        writer.writerows(problem2_cv_rows)

    rusty_baseline = next(row for row in problem2_detail_rows if row[0] == "无交互" and row[1] == rusty)
    rusty_threshold = next(row for row in problem2_detail_rows if row[0] == "门槛交互" and row[1] == rusty)
    problem2_rusty_path = out / "问题2_锈蚀钢筋改善汇总.csv"
    with open_output_text(problem2_rusty_path, encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["对比项", "RMSE(mT)", "MAE(mT)", "MAPE"])
        writer.writerow(["无交互", rusty_baseline[2], rusty_baseline[3], rusty_baseline[4]])
        writer.writerow(["门槛交互", rusty_threshold[2], rusty_threshold[3], rusty_threshold[4]])
        writer.writerow([
            "RMSE改善比例",
            round((rusty_baseline[2] - rusty_threshold[2]) / rusty_baseline[2], 6),
            round((rusty_baseline[3] - rusty_threshold[3]) / rusty_baseline[3], 6),
            round((rusty_baseline[4] - rusty_threshold[4]) / rusty_baseline[4], 6),
        ])

    problem2_rusty_2290_path = out / "问题2_锈蚀钢筋_22_90天汇总.csv"
    with open_output_text(problem2_rusty_2290_path, encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["模型版本", "RMSE(mT)", "MAE(mT)", "MAPE"])
        for variant_name, mode, _ in variant_specs:
            sq, ab, ape = [], [], []
            x_rows_v, y_v = [], []
            for row in observations:
                if row["day"] == 0 or row["type"] != rusty or row["day"] < 22:
                    continue
                x_rows_v.append(features_ablation(row["type"], row["day"], mode))
                y_v.append(math.log(m0[(row["type"], row["id"])] / row["mag"]))
            if len(x_rows_v) < 5:
                writer.writerow([variant_name, "", "", ""])
                continue
            fit_v = fit_ridge(x_rows_v, y_v)
            held_rows = [r for r in observations if r["day"] >= 22 and r["type"] == rusty and r["day"] > 0]
            for row, yhat in zip(held_rows, fit_v["yhat"]):
                pred = m0[(row["type"], row["id"])] * math.exp(-yhat)
                err = row["mag"] - pred
                sq.append(err * err)
                ab.append(abs(err))
                ape.append(abs(err) / row["mag"])
            writer.writerow([variant_name, round((sum(sq) / len(sq)) ** 0.5, 6), round(sum(ab) / len(ab), 6), round(sum(ape) / len(ape), 6)])

    problem2_window_path = out / "问题2_高风险窗口对比.csv"
    with open_output_text(problem2_window_path, encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["模型版本", "高风险窗口日数", "RMSE(mT)", "MAE(mT)", "MAPE"])
        high_risk_days = [d for d in range(1, 91) if weather[d]["temp"] >= temp_mean and weather[d]["hum"] >= hum_mean]
        for variant_name, mode, _ in variant_specs:
            sq, ab, ape = [], [], []
            x_rows_v, y_v = [], []
            for row in observations:
                if row["day"] == 0 or row["day"] not in high_risk_days:
                    continue
                x_rows_v.append(features_ablation(row["type"], row["day"], mode))
                y_v.append(math.log(m0[(row["type"], row["id"])] / row["mag"]))
            if len(x_rows_v) < 5:
                writer.writerow([variant_name, len(high_risk_days), "", "", ""])
                continue
            fit_v = fit_ridge(x_rows_v, y_v)
            held_rows = [r for r in observations if r["day"] in high_risk_days and r["day"] > 0]
            for row, yhat in zip(held_rows, fit_v["yhat"]):
                pred = m0[(row["type"], row["id"])] * math.exp(-yhat)
                err = row["mag"] - pred
                sq.append(err * err)
                ab.append(abs(err))
                ape.append(abs(err) / row["mag"])
            writer.writerow([variant_name, len(high_risk_days), round((sum(sq) / len(sq)) ** 0.5, 6), round(sum(ab) / len(ab), 6), round(sum(ape) / len(ape), 6)])

    problem2_variant_specs = [
        ("无交互", "none"),
        ("对称交互", "symmetric"),
        ("门槛交互", "threshold"),
    ]

    def fit_problem2_variant(mode: str, use_ridge: bool = False, subset_filter=None):
        rows = [row for row in observed_rows_day if subset_filter is None or subset_filter(row)]
        x_rows_variant = [features_problem2(row["type"], row["day"], mode) for row in rows]
        y_variant = [math.log(m0[(row["type"], row["id"])] / row["mag"]) for row in rows]
        fit_variant = fit_ridge(x_rows_variant, y_variant, lam=0.1) if use_ridge else fit_ols(x_rows_variant, y_variant)
        overall_metrics, metrics_by_type_variant = summarize_prediction_errors(
            rows, fit_variant["yhat"], m0, types_order
        )
        return {
            "rows": rows,
            "x_rows": x_rows_variant,
            "y": y_variant,
            "fit": fit_variant,
            "overall": overall_metrics,
            "by_type": metrics_by_type_variant,
        }

    def cross_validate_problem2_variant(mode: str):
        cv_sq, cv_ab, cv_ape = [], [], []
        for sample_type in types_order:
            sample_ids = sorted({row["id"] for row in observations if row["type"] == sample_type and row["day"] == 0})
            for sample_id in sample_ids:
                train_rows = [
                    row for row in observed_rows_day
                    if not (row["type"] == sample_type and row["id"] == sample_id)
                ]
                x_train = [features_problem2(row["type"], row["day"], mode) for row in train_rows]
                y_train = [math.log(m0[(row["type"], row["id"])] / row["mag"]) for row in train_rows]
                fit_cv = fit_ols(x_train, y_train)
                held_rows = [
                    row for row in observed_rows_day
                    if row["type"] == sample_type and row["id"] == sample_id
                ]
                yhat_cv = [
                    sum(beta * value for beta, value in zip(fit_cv["beta"], features_problem2(row["type"], row["day"], mode)))
                    for row in held_rows
                ]
                for row, pred_log in zip(held_rows, yhat_cv):
                    pred = m0[(row["type"], row["id"])] * math.exp(-pred_log)
                    err = row["mag"] - pred
                    cv_sq.append(err * err)
                    cv_ab.append(abs(err))
                    cv_ape.append(abs(err) / row["mag"])
        return {
            "rmse": (sum(cv_sq) / len(cv_sq)) ** 0.5 if cv_sq else float("nan"),
            "mae": sum(cv_ab) / len(cv_ab) if cv_ab else float("nan"),
            "mape": sum(cv_ape) / len(cv_ape) if cv_ape else float("nan"),
        }

    problem2_summary_rows = []
    problem2_detail_rows = []
    problem2_cv_rows = []
    variant_result_map = {}
    for variant_name, mode in problem2_variant_specs:
        variant_result = fit_problem2_variant(mode)
        variant_result_map[mode] = variant_result
        fit_variant = variant_result["fit"]
        by_type_variant = variant_result["by_type"]
        key_beta = fit_variant["beta"][-1] if mode != "none" else ""
        key_p = normal_p_value(fit_variant["beta"][-1] / fit_variant["se"][-1]) if mode != "none" else ""
        problem2_summary_rows.append(
            [
                variant_name,
                round(fit_variant["r2"], 6),
                round(fit_variant["rmse"], 6),
                round(variant_result["overall"]["rmse_m"], 6),
                round(variant_result["overall"]["mae_m"], 6),
                round(variant_result["overall"]["mape_m"], 6),
                round(key_beta, 6) if key_beta != "" else "",
                round(key_p, 6) if key_p != "" else "",
            ]
        )
        for sample_type in types_order:
            item = by_type_variant[sample_type]
            problem2_detail_rows.append(
                [variant_name, sample_type, round(item["rmse"], 6), round(item["mae"], 6), round(item["mape"], 6)]
            )
        cv_metrics = cross_validate_problem2_variant(mode)
        problem2_cv_rows.append(
            [variant_name, round(cv_metrics["rmse"], 6), round(cv_metrics["mae"], 6), round(cv_metrics["mape"], 6)]
        )

    problem2_summary_path = out / "问题2_消融分析.csv"
    write_csv(
        problem2_summary_path,
        ["模型版本", "R2_log", "RMSE_log", "RMSE_m", "MAE_m", "MAPE_m", "末项系数", "末项p值"],
        problem2_summary_rows,
    )

    problem2_detail_path = out / "问题2_类型误差对比.csv"
    write_csv(problem2_detail_path, ["模型版本", "样品类型", "RMSE(mT)", "MAE(mT)", "MAPE"], problem2_detail_rows)

    problem2_cv_path = out / "问题2_交叉验证.csv"
    write_csv(problem2_cv_path, ["模型版本", "CV_RMSE(mT)", "CV_MAE(mT)", "CV_MAPE"], problem2_cv_rows)

    rusty_baseline = next(row for row in problem2_detail_rows if row[0] == "无交互" and row[1] == rusty)
    rusty_threshold = next(row for row in problem2_detail_rows if row[0] == "门槛交互" and row[1] == rusty)
    problem2_rusty_path = out / "问题2_锈蚀钢筋改善汇总.csv"
    write_csv(
        problem2_rusty_path,
        ["对比项", "RMSE(mT)", "MAE(mT)", "MAPE"],
        [
            ["无交互", rusty_baseline[2], rusty_baseline[3], rusty_baseline[4]],
            ["门槛交互", rusty_threshold[2], rusty_threshold[3], rusty_threshold[4]],
            [
                "RMSE改善比例",
                round((rusty_baseline[2] - rusty_threshold[2]) / rusty_baseline[2], 6),
                round((rusty_baseline[3] - rusty_threshold[3]) / rusty_baseline[3], 6),
                round((rusty_baseline[4] - rusty_threshold[4]) / rusty_baseline[4], 6),
            ],
        ],
    )

    problem2_rusty_2290_path = out / "问题2_锈蚀钢筋_22_90天汇总.csv"
    rusty_window_rows = []
    for variant_name, mode in problem2_variant_specs:
        variant_result = fit_problem2_variant(
            mode,
            use_ridge=True,
            subset_filter=lambda row: row["type"] == rusty and row["day"] >= 22,
        )
        rusty_window_rows.append(
            [
                variant_name,
                round(variant_result["overall"]["rmse_m"], 6),
                round(variant_result["overall"]["mae_m"], 6),
                round(variant_result["overall"]["mape_m"], 6),
            ]
        )
    write_csv(problem2_rusty_2290_path, ["模型版本", "RMSE(mT)", "MAE(mT)", "MAPE"], rusty_window_rows)

    high_risk_days = [d for d in range(1, 91) if weather[d]["temp"] >= temp_mean and weather[d]["hum"] >= hum_mean]
    problem2_window_path = out / "问题2_高风险窗口对比.csv"
    high_risk_rows = []
    for variant_name, mode in problem2_variant_specs:
        variant_result = fit_problem2_variant(
            mode,
            use_ridge=True,
            subset_filter=lambda row, days=set(high_risk_days): row["day"] in days,
        )
        high_risk_rows.append(
            [
                variant_name,
                len(high_risk_days),
                round(variant_result["overall"]["rmse_m"], 6),
                round(variant_result["overall"]["mae_m"], 6),
                round(variant_result["overall"]["mape_m"], 6),
            ]
        )
    write_csv(problem2_window_path, ["模型版本", "高风险窗口日数", "RMSE(mT)", "MAE(mT)", "MAPE"], high_risk_rows)

    # Prediction table for days 23-29.
    pred_path = out / "问题1_第23-29天预测.csv"
    with open_output_text(pred_path, encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["测量天数"] + [f"{t}预测剩磁(mT)" for t in types_order])
        for day in range(23, 30):
            row = [day]
            for sample_type in types_order:
                row.append(round(initial_stats[sample_type]["mean"] * survival(sample_type, day), 4))
            writer.writerow(row)

    # Dynamic threshold table and type-mean remanence table.
    threshold_headers = [
        "序号",
        "年",
        "月",
        "日",
        "天气",
        "气温(℃)",
        "湿度(%)",
        f"{nail}动态阈值(mT)",
        f"{clip}动态阈值(mT)",
        f"{ordinary}动态阈值(mT)",
        f"{rusty}动态阈值(mT)",
        f"{nail}预测剩磁均值(mT)",
        f"{clip}预测剩磁均值(mT)",
        f"{ordinary}预测剩磁均值(mT)",
        f"{rusty}预测剩磁均值(mT)",
    ]
    threshold_rows = []
    for day in range(1, 91):
        weather_day = weather[day]
        row = [
            day,
            weather_day["year"],
            weather_day["month"],
            weather_day["date"],
            weather_day["weather"],
            weather_day["temp"],
            weather_day["hum"],
        ]
        for sample_type in types_order:
            row.append(threshold_profile_by_type[sample_type][day - 1])
        for sample_type in types_order:
            row.append(theta_obs(sample_type, day, use_mean=True))
        threshold_rows.append(row)
    threshold_xlsx = out / "问题3_动态阈值修正表.xlsx"
    write_simple_xlsx(threshold_xlsx, threshold_headers, threshold_rows)

    problem3_key_path = out / "问题3_阈值关键点.csv"
    key_days = [1, 7, 15, 22, 30, 58, 90]
    problem3_key_rows = []
    for day in key_days:
        weather_day = weather[day]
        row = [
            day,
            weather_day["weather"],
            weather_day["temp"],
            weather_day["hum"],
        ]
        for sample_type in types_order:
            row.append(threshold_profile_by_type[sample_type][day - 1])
        for sample_type in types_order:
            row.append(theta_obs(sample_type, day, use_mean=True))
        problem3_key_rows.append(row)
    write_csv(
        problem3_key_path,
        [
            "天数",
            "天气",
            "气温(℃)",
            "湿度(%)",
            f"{nail}动态阈值(mT)",
            f"{clip}动态阈值(mT)",
            f"{ordinary}动态阈值(mT)",
            f"{rusty}动态阈值(mT)",
            f"{nail}类型均值预测剩磁(mT)",
            f"{clip}类型均值预测剩磁(mT)",
            f"{ordinary}类型均值预测剩磁(mT)",
            f"{rusty}类型均值预测剩磁(mT)",
        ],
        problem3_key_rows,
    )

    problem3_static_path = out / "问题3_固定阈值对照.csv"
    problem3_static_rows = []
    for day in key_days:
        weather_day = weather[day]
        row = [day, weather_day["weather"], weather_day["temp"], weather_day["hum"]]
        for sample_type in types_order:
            row.append(theta_static(sample_type))
        for sample_type in types_order:
            row.append(threshold_profile_by_type[sample_type][day - 1])
        problem3_static_rows.append(row)
    write_csv(
        problem3_static_path,
        [
            "天数",
            "天气",
            "气温(℃)",
            "湿度(%)",
            f"{nail}固定阈值(mT)",
            f"{clip}固定阈值(mT)",
            f"{ordinary}固定阈值(mT)",
            f"{rusty}固定阈值(mT)",
            f"{nail}映射阈值(mT)",
            f"{clip}映射阈值(mT)",
            f"{ordinary}映射阈值(mT)",
            f"{rusty}映射阈值(mT)",
        ],
        problem3_static_rows,
    )

    problem3_mono_path = out / "问题3_单调性校验.csv"
    problem3_mono_rows = []
    threshold_by_type = {
        sample_type: [row[7 + idx] for row in threshold_rows] for idx, sample_type in enumerate(types_order)
    }
    for sample_type, values in threshold_by_type.items():
        increases = sum(1 for a, b in zip(values, values[1:]) if b > a + 1e-12)
        problem3_mono_rows.append(
            [
                sample_type,
                "是" if increases == 0 else "否",
                increases,
                round(values[0], 6),
                round(values[-1], 6),
            ]
        )
    write_csv(problem3_mono_path, ["样品类型", "是否单调不增", "上升次数", "首日阈值", "末日阈值"], problem3_mono_rows)

    problem3_boot_path = out / "问题3_阈值区间.csv"
    problem3_boot_rows = []
    cluster_rows_problem3 = {}
    for row in observations:
        if row["day"] == 0:
            continue
        cluster_rows_problem3.setdefault((row["type"], row["id"]), []).append(row)
    cluster_keys_problem3 = list(cluster_rows_problem3.keys())
    rng_problem3 = random.Random(20260513)
    boot_theta_map = {(day, sample_type): [] for day in key_days for sample_type in types_order}
    bootstrap_reps_problem3 = 200
    for _ in range(bootstrap_reps_problem3):
        sampled_keys = [rng_problem3.choice(cluster_keys_problem3) for _ in cluster_keys_problem3]
        boot_rows = []
        for key in sampled_keys:
            boot_rows.extend(cluster_rows_problem3[key])
        boot_x_rows, boot_y = [], []
        for row in boot_rows:
            boot_x_rows.append(features(row["type"], row["day"]))
            boot_y.append(math.log(m0[(row["type"], row["id"])] / row["mag"]))
        try:
            boot_model = fit_ols(boot_x_rows, boot_y)
        except ValueError:
            continue

        def boot_survival(sample_type: str, day: int) -> float:
            x = features(sample_type, day)
            yhat = sum(bi * xi for bi, xi in zip(boot_model["beta"], x))
            return math.exp(-yhat)

        for day in key_days:
            for sample_type in types_order:
                theta0 = 1.0 if sample_type in (nail, clip) else 1.5
                boot_theta_map[(day, sample_type)].append(theta0 * boot_survival(sample_type, day))

    for day in key_days:
        weather_day = weather[day]
        for sample_type in types_order:
            vals = boot_theta_map[(day, sample_type)]
            if vals:
                lower = percentile(vals, 0.025)
                upper = percentile(vals, 0.975)
                point = threshold_profile_by_type[sample_type][day - 1]
            else:
                lower = upper = point = float("nan")
            problem3_boot_rows.append(
                [
                    day,
                    weather_day["weather"],
                    sample_type,
                    round(point, 6) if not math.isnan(point) else "",
                    round(lower, 6) if not math.isnan(lower) else "",
                    round(upper, 6) if not math.isnan(upper) else "",
                    len(vals),
                ]
            )
    write_csv(problem3_boot_path, ["天数", "天气", "样品类型", "点估计阈值", "下界", "上界", "Bootstrap次数"], problem3_boot_rows)

    # Coefficient and metric table.
    params_path = out / "模型参数与检验.csv"
    with open_output_text(params_path, encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["变量", "系数", "标准误", "t值", "近似p值"])
        for label, beta, se in zip(labels, model["beta"], model["se"]):
            t_value = beta / se
            p_value = normal_p_value(t_value)
            writer.writerow([label, round(beta, 6), round(se, 6), round(t_value, 3), "<1e-6" if p_value < 1e-6 else round(p_value, 6)])
        writer.writerow([])
        writer.writerow(["指标", "数值"])
        for key, value in metrics.items():
            writer.writerow([key, round(value, 6)])
        writer.writerow([])
        writer.writerow(["样品类型", "RMSE(mT)", "MAE(mT)", "MAPE"])
        for sample_type in types_order:
            item = metrics_by_type[sample_type]
            writer.writerow([sample_type, round(item["rmse"], 6), round(item["mae"], 6), round(item["mape"], 6)])

    problem2_main_path = out / "问题2_主效应与交互效应.csv"
    with open_output_text(problem2_main_path, encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["变量", "系数", "标准误", "t值", "近似p值", "含义"])
        meaning = {
            "截距": "基准衰减截距",
            "小号铁钉": "相对普通钢筋的类型效应",
            "小号铁夹": "相对普通钢筋的类型效应",
            "锈蚀钢筋": "相对普通钢筋的类型效应",
            "时间_z": "时间主效应",
            "时间_z^2": "时间非线性修正",
            "累计温度_z": "温度累计主效应",
            "累计湿度_z": "湿度累计主效应",
            "累计温湿交互_z": "温湿交互主效应",
            "锈蚀钢筋×时间_z": "锈蚀钢筋时间交互",
            "小号铁钉×时间_z": "小号铁钉时间交互",
            "小号铁夹×时间_z": "小号铁夹时间交互",
        }
        for label, beta, se in zip(labels, model["beta"], model["se"]):
            t_value = beta / se
            p_value = normal_p_value(t_value)
            writer.writerow([label, round(beta, 6), round(se, 6), round(t_value, 3), "<1e-6" if p_value < 1e-6 else round(p_value, 6), meaning.get(label, "")])

    problem2_main_path = out / "问题2_主效应与交互效应.csv"
    problem2_main_rows = []
    problem2_meaning = {
        "截距": "基准衰减截距",
        "小号铁钉": "相对普通钢筋的类型效应",
        "小号铁夹": "相对普通钢筋的类型效应",
        "锈蚀钢筋": "相对普通钢筋的类型效应",
        "时间_z": "时间主效应",
        "时间_z^2": "时间非线性修正",
        "累计温度_z": "去时间趋势后的累计温度主效应",
        "累计湿度_z": "去时间趋势后的累计湿度主效应",
        "锈蚀钢筋×时间_z": "锈蚀钢筋时间交互",
        "小号铁钉×时间_z": "小号铁钉时间交互",
        "小号铁夹×时间_z": "小号铁夹时间交互",
        "累计温湿交互_z": "去时间趋势后的温湿交互主效应",
    }
    for label, beta, se in zip(problem2_labels, problem2_model["beta"], problem2_cluster_se):
        t_value = beta / se
        p_value = normal_p_value(t_value)
        problem2_main_rows.append(
            [
                label,
                round(beta, 6),
                round(se, 6),
                round(t_value, 3),
                "<1e-6" if p_value < 1e-6 else round(p_value, 6),
                problem2_meaning.get(label, ""),
            ]
        )
    write_csv(problem2_main_path, ["变量", "系数", "聚类稳健标准误", "t值", "近似p值", "含义"], problem2_main_rows)

    # Problem 4: cluster bootstrap and three-region decision table.
    cluster_rows = {}
    for row in observations:
        if row["day"] == 0:
            continue
        cluster_rows.setdefault((row["type"], row["id"]), []).append(row)

    target_rows = [row for row in observations if row["day"] > 0]
    target_boot_values = [[] for _ in target_rows]
    cluster_keys = list(cluster_rows.keys())
    rng = random.Random(20260513)
    bootstrap_reps = 1000

    for _ in range(bootstrap_reps):
        sampled_keys = [rng.choice(cluster_keys) for _ in cluster_keys]
        boot_rows = []
        for key in sampled_keys:
            boot_rows.extend(cluster_rows[key])

        boot_x_rows, boot_y = [], []
        for row in boot_rows:
            boot_x_rows.append(features(row["type"], row["day"]))
            boot_y.append(math.log(m0[(row["type"], row["id"])] / row["mag"]))

        try:
            boot_model = fit_ols(boot_x_rows, boot_y)
        except ValueError:
            continue

        def boot_predict_log_decay(sample_type: str, day: int) -> float:
            x = features(sample_type, day)
            return sum(bi * xi for bi, xi in zip(boot_model["beta"], x))

        def boot_survival(sample_type: str, day: int) -> float:
            return math.exp(-boot_predict_log_decay(sample_type, day))

        for idx, row in enumerate(target_rows):
            m0_hat_boot = row["mag"] / boot_survival(row["type"], row["day"])
            target_boot_values[idx].append(m0_hat_boot)

    problem4_path = out / "问题4_三区域判定结果.csv"
    with open_output_text(problem4_path, encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(
            [
                "样品类型",
                "样品编号",
                "测量天数",
                "观测剩磁(mT)",
                "保留率R",
                "动态阈值(mT)",
                "反推初值点估计(mT)",
                "反推初值下界(mT)",
                "反推初值上界(mT)",
                "区间相对宽度",
                "三区域判定",
                "辅助概率",
                "高风险提示",
            ]
        )
        for row, boot_vals in zip(target_rows, target_boot_values):
            r_hat = survival(row["type"], row["day"])
            theta_t = (1.0 if row["type"] in (nail, clip) else 1.5) * r_hat
            m0_hat = row["mag"] / r_hat
            theta0 = 1.0 if row["type"] in (nail, clip) else 1.5
            if boot_vals:
                lower = percentile(boot_vals, 0.025)
                upper = percentile(boot_vals, 0.975)
                relative_width = (upper - lower) / m0_hat if m0_hat else float("inf")
                support_prob = sum(1 for v in boot_vals if v >= theta0) / len(boot_vals)
            else:
                lower = upper = float("nan")
                relative_width = float("nan")
                support_prob = float("nan")
            gray_zone = (
                math.isnan(relative_width)
                or relative_width > 1.0
                or r_hat < 0.05
            )
            if gray_zone:
                decision = "灰区/证据不足"
            elif lower >= theta0:
                decision = "支持曾遭雷击"
            elif upper < theta0:
                decision = "不支持曾遭雷击"
            else:
                decision = "灰区/证据不足"
            warning_flag = "是" if (not math.isnan(relative_width) and (relative_width > 1.0 or r_hat < 0.05)) else "否"
            writer.writerow(
                [
                    row["type"],
                    row["id"],
                    row["day"],
                    round(row["mag"], 6),
                    round(r_hat, 6),
                    round(theta_t, 6),
                    round(m0_hat, 6),
                    round(lower, 6) if not math.isnan(lower) else "",
                    round(upper, 6) if not math.isnan(upper) else "",
                    round(relative_width, 6) if not math.isnan(relative_width) else "",
                    decision,
                    round(support_prob, 6) if not math.isnan(support_prob) else "",
                    warning_flag,
                ]
            )

    # Main report.
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(10.5)

        doc.add_heading("B题 剩磁法雷击判定预测完整求解报告", level=1)
        doc.add_paragraph(
            "本报告读取题目 PDF、附件1模拟实验数据、附件2逐日天气数据，并以已有 Word 初稿中的"
            "“环境驱动指数衰减”思路为基础，完成参数估计、因素分析、动态阈值修正和雷击判定流程设计。"
        )

        doc.add_heading("一、数据预处理", level=2)
        doc.add_paragraph(
            "附件1共1380条观测，四类样品分别为小号铁钉、小号铁夹、普通钢筋和锈蚀钢筋。"
            "以“样品类型+编号”为唯一样品标识，将第0天剩磁作为该样品初始剩磁 M0，"
            "并构造相对剩磁衰减量 Y=ln(M0/Mt)。附件2提供第1至90天逐日温度和相对湿度。"
        )
        doc.add_paragraph(
            "缺失检查显示本题附件可直接匹配；若实际应用中存在缺失值，剩磁序列采用同一样品时间序列插值，"
            "天气变量采用相邻日期线性插值。异常值可按同类型同天数箱线图或3σ准则筛查。"
        )

        doc.add_heading("二、问题一：动态衰减模型", level=2)
        doc.add_paragraph(
            "设 M_sj(t) 为类型 s、第 j 个样品在第 t 天的剩磁值，M_sj(0) 为其初始剩磁。"
            "采用指数衰减框架：M_sj(t)=M_sj(0) exp[-Y_s(t)]。"
        )
        doc.add_paragraph(
            "令 CT(t)=Σ(Td-Tbar)，CH(t)=Σ(Hd-Hbar)，CTH(t)=Σ[(Td-Tbar)(Hd-Hbar)]，"
            "并将时间和三个累计环境量标准化。以普通钢筋为基准类型，拟合模型："
        )
        doc.add_paragraph(
            "Y=β0+βN·N+βC·C+βR·R+β1·zt+β2·zt²+βT·CTz+βH·CHz+βTH·CTHz"
            "+βRt·Rzt+βNt·Nzt+βCt·Czt+ε。"
        )
        doc.add_paragraph(
            f"模型在对数衰减量上的 R²={metrics['r2_log']:.4f}，RMSE={metrics['rmse_log']:.4f}；"
            f"换算到剩磁值，RMSE={metrics['rmse_m']:.4f} mT，MAE={metrics['mae_m']:.4f} mT，"
            f"MAPE={metrics['mape_m']:.2%}。"
        )

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "样品类型"
        hdr[1].text = "初始均值(mT)"
        hdr[2].text = "标准差"
        hdr[3].text = "最小值"
        hdr[4].text = "最大值"
        for sample_type in types_order:
            item = initial_stats[sample_type]
            cells = table.add_row().cells
            cells[0].text = sample_type
            cells[1].text = f"{item['mean']:.4f}"
            cells[2].text = f"{item['sd']:.4f}"
            cells[3].text = f"{item['min']:.4f}"
            cells[4].text = f"{item['max']:.4f}"

        doc.add_paragraph("第23至29天各类型预测剩磁均值如下：")
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "天数"
        for idx, sample_type in enumerate(types_order, start=1):
            hdr[idx].text = sample_type
        for day in range(23, 30):
            cells = table.add_row().cells
            cells[0].text = str(day)
            for idx, sample_type in enumerate(types_order, start=1):
                pred = initial_stats[sample_type]["mean"] * survival(sample_type, day)
                cells[idx].text = f"{pred:.4f}"

        doc.add_heading("三、问题二：关键影响因素分析", level=2)
        doc.add_paragraph(
            "温度×湿度交互项系数为正且显著，βTH="
            f"{problem2_model['beta'][11]:.4f}，t={problem2_model['beta'][11]/problem2_cluster_se[11]:.2f}，p<1e-6。"
            "这说明在控制时间和样品类型后，高温高湿组合会产生额外的耦合衰减贡献。"
        )
        doc.add_paragraph(
            "锈蚀钢筋主效应和锈蚀×时间项均显著，说明锈蚀不仅提高累计衰减水平，还会随检测滞后时间"
            "产生累积放大。普通钢筋第30天保留率约为"
            f"{survival(ordinary, 30):.2%}，锈蚀钢筋为{survival(rusty, 30):.2%}；"
            f"第90天分别为{survival(ordinary, 90):.3%}和{survival(rusty, 90):.3%}。"
        )
        doc.add_paragraph(
            "样品类型差异显著：普通钢筋稳定性最好，小号铁钉和小号铁夹因尺寸小、形状影响而衰减更快，"
            "锈蚀钢筋因锈蚀破坏局部磁畴结构，长期衰减最快。"
        )

        doc.add_heading("四、问题三：动态阈值修正", level=2)
        doc.add_paragraph(
            "国标静态阈值中，小尺寸铁件为1.0 mT，大尺寸铁件为1.5 mT。"
            "根据衰减模型，动态阈值定义为 θs(t)=θs(0) exp[-Ys(t)]。"
            "已生成1至90天动态阈值 Excel 表：问题3_动态阈值修正表.xlsx。"
        )
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "天数"
        for idx, sample_type in enumerate(types_order, start=1):
            hdr[idx].text = f"{sample_type}动态阈值"
        for day in [1, 7, 15, 22, 30, 58, 90]:
            cells = table.add_row().cells
            cells[0].text = str(day)
            for idx, sample_type in enumerate(types_order, start=1):
                theta0 = 1.0 if sample_type in (nail, clip) else 1.5
                cells[idx].text = f"{theta0 * survival(sample_type, day):.4f}"

        doc.add_heading("五、问题四：雷击判定支持流程", level=2)
        doc.add_paragraph(
            "流程为：1）输入样品类型、锈蚀等级、检测延迟天数、实测剩磁及天气信息；"
            "2）进行缺失值和异常值处理，统一单位并匹配天气序列；"
            "3）调用衰减模型计算当前剩磁预测和动态阈值；"
            "4）将实测剩磁与动态阈值比较；"
            "5）根据对数尺度误差输出置信水平。"
        )
        doc.add_paragraph(
            f"置信度可取 Z=[ln(Mobs)-ln(θs(t))]/{metrics['sigma_log']:.4f}，P=Φ(Z)。"
            "建议等级为：P≥0.90高度支持雷击，0.70≤P<0.90较支持雷击，"
            "0.40≤P<0.70证据不足，P<0.40不支持雷击。"
        )

        report_path = out / "B题_完整求解报告.docx"
        doc.save(report_path)
    except Exception as exc:
        report_path = out / "B题_完整求解报告.md"
        with open_output_text(report_path, encoding="utf-8") as f:
            f.write("# B题完整求解报告\n\n")
            f.write("## 生成状态\n")
            f.write(f"- Word报告生成失败，已改为Markdown兜底：{exc}\n\n")
            f.write("## 核心输出\n")
            f.write(f"- 问题1预测表：{pred_path.name}\n")
            f.write(f"- 问题1地基检验：{problem1_path.name}\n")
            f.write(f"- 问题1 LOSO交叉验证：{problem1_cv_path.name}\n")
            f.write(f"- 问题3动态阈值表：{threshold_xlsx.name}\n")
            f.write(f"- 问题3阈值关键点：{problem3_key_path.name}\n")
            f.write(f"- 问题3固定阈值对照：{problem3_static_path.name}\n")
            f.write(f"- 问题3单调性校验：{problem3_mono_path.name}\n")
            f.write(f"- 问题3阈值区间：{problem3_boot_path.name}\n")
            f.write(f"- 模型参数与检验：{params_path.name}\n")
            f.write(f"- 问题2主效应与交互效应：{problem2_main_path.name}\n")
            f.write(f"- 问题2消融分析：{problem2_summary_path.name}\n")
            f.write(f"- 问题2类型误差对比：{problem2_detail_path.name}\n")
            f.write(f"- 问题2交叉验证：{problem2_cv_path.name}\n")
            f.write(f"- 问题2锈蚀钢筋改善汇总：{problem2_rusty_path.name}\n")
            f.write(f"- 问题2锈蚀钢筋22-90天汇总：{problem2_rusty_2290_path.name}\n")
            f.write(f"- 问题2高风险窗口对比：{problem2_window_path.name}\n")
            f.write(f"- 问题2共线性诊断：{problem2_collinearity_path.name}\n")
            f.write(f"- 问题2VIF诊断：{problem2_vif_path.name}\n")
            f.write(f"- 问题2岭回归稳健性：{problem2_ridge_path.name}\n")
            f.write(f"- 问题2收缩模型对比：{problem2_shrink_path.name}\n")
            f.write(f"- 问题4三区域判定：{problem4_path.name}\n\n")
            f.write("## 当前模型口径\n")
            f.write("- 个体初值锚定 + 类型共享衰减函数\n")
            f.write("- 动态阈值映射：Theta_s(t)=Theta_s^{GB}·R_s(t)\n")
            f.write("- 反推初值：M0_hat = Mobs / R_s(t)\n")
            f.write("- 问题4：Bootstrap区间 + 三区域判定\n")
        print(f"Report generation fallback to markdown: {exc}")

    print("Generated outputs:")
    problem1_path = out / "问题1_地基检验结果.csv"
    problem1_cv_path = out / "问题1_LOSO交叉验证.csv"
    problem2_main_path = out / "问题2_主效应与交互效应.csv"
    problem2_summary_path = out / "问题2_消融分析.csv"
    problem2_detail_path = out / "问题2_类型误差对比.csv"
    problem2_cv_path = out / "问题2_交叉验证.csv"
    problem2_rusty_path = out / "问题2_锈蚀钢筋改善汇总.csv"
    problem2_rusty_2290_path = out / "问题2_锈蚀钢筋_22_90天汇总.csv"
    problem2_window_path = out / "问题2_高风险窗口对比.csv"
    problem4_path = out / "问题4_三区域判定结果.csv"
    for path in [pred_path, problem1_path, problem1_cv_path, threshold_xlsx, problem3_key_path, problem3_static_path, problem3_mono_path, problem3_boot_path, params_path, problem2_main_path, problem2_summary_path, problem2_detail_path, problem2_cv_path, problem2_rusty_path, problem2_rusty_2290_path, problem2_window_path, problem2_collinearity_path, problem2_vif_path, problem2_ridge_path, problem2_shrink_path, problem4_path, report_path]:
        if path:
            print(path)


if __name__ == "__main__":
    main()
